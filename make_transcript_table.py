from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_font(cell, text, bold=False, size=11, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if color:
        run.font.color.rgb = RGBColor(*color)

def set_cell_left(cell, text, bold=False, size=11):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'

# ---- Data ----
spring_courses = [
    ("CSC630", "Advanced Data Mining",                    "3", "3.33"),
    ("CSC631", "Advanced Topics in Computer Networks",    "3", "4"),
    ("CSC652", "Advanced Topics in Computer Vision",      "3", "4"),
    ("CSC683", "Advanced Algorithm Analysis",             "3", "4"),
]
spring_cgpa = "3.83"

fall_courses = [
    ("CSC512", "Theory of Automata – II",        "3", "3.66"),
    ("CSC602", "Advanced Operating Systems",     "3", "4"),
    ("CSC604", "Advanced Computer Architecture", "3", "4"),
    ("CSC607", "Research Methodology",           "1", "3.66"),
]
fall_cgpa = "3.85"

HEADER_BG   = "D9D9D9"   # light grey for column headers
SECTION_BG  = "F2F2F2"   # slightly lighter for semester label rows
WHITE       = "FFFFFF"

total_rows = (
    1 +                     # col headers
    1 +                     # Spring label
    len(spring_courses) +
    1 +                     # Spring CGPA
    1 +                     # Fall label
    len(fall_courses) +
    1                       # Fall CGPA
)

table = doc.add_table(rows=total_rows, cols=4)
table.style = 'Table Grid'

# column widths
widths = [Inches(0.5), Inches(3.2), Inches(1.2), Inches(2.1)]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

row_idx = 0

# ---- HEADER ROW ----
hdr = table.rows[row_idx].cells
set_cell_bg(hdr[0], HEADER_BG)
set_cell_bg(hdr[1], HEADER_BG)
set_cell_bg(hdr[2], HEADER_BG)
set_cell_bg(hdr[3], HEADER_BG)
set_cell_font(hdr[0], "",                        bold=True)
set_cell_left(hdr[1], "Course Code and Title",   bold=True)
set_cell_font(hdr[2], "Credit Hours",            bold=True)
set_cell_font(hdr[3], "Grade Points / Semester", bold=True)
row_idx += 1

# ---- SPRING SEMESTER LABEL ----
sem_row = table.rows[row_idx].cells
for c in sem_row:
    set_cell_bg(c, SECTION_BG)
sem_row[0].merge(sem_row[3])
set_cell_left(sem_row[0], "Result Semester: Spring 2025", bold=True, size=11)
row_idx += 1

# ---- SPRING COURSES ----
for i, (code, title, credit, gp) in enumerate(spring_courses, start=1):
    cells = table.rows[row_idx].cells
    set_cell_font(cells[0], str(i) + ".",  bold=False)
    set_cell_left(cells[1], f"({code})  {title}")
    set_cell_font(cells[2], credit)
    set_cell_font(cells[3], gp)
    row_idx += 1

# ---- SPRING CGPA ROW ----
cgpa_row = table.rows[row_idx].cells
for c in cgpa_row:
    set_cell_bg(c, SECTION_BG)
set_cell_left(cgpa_row[0], f"CGPA: {spring_cgpa}", bold=True)
cgpa_row[1].merge(cgpa_row[3])
set_cell_left(cgpa_row[1], "Scholastic Status: GAS", bold=True)
row_idx += 1

# ---- FALL SEMESTER LABEL ----
sem_row2 = table.rows[row_idx].cells
for c in sem_row2:
    set_cell_bg(c, SECTION_BG)
sem_row2[0].merge(sem_row2[3])
set_cell_left(sem_row2[0], "Result Semester: Fall 2025", bold=True, size=11)
row_idx += 1

# ---- FALL COURSES ----
for i, (code, title, credit, gp) in enumerate(fall_courses, start=1):
    cells = table.rows[row_idx].cells
    set_cell_font(cells[0], str(i) + ".",  bold=False)
    set_cell_left(cells[1], f"({code})  {title}")
    set_cell_font(cells[2], credit)
    set_cell_font(cells[3], gp)
    row_idx += 1

# ---- FALL CGPA ROW ----
cgpa_row2 = table.rows[row_idx].cells
for c in cgpa_row2:
    set_cell_bg(c, SECTION_BG)
set_cell_left(cgpa_row2[0], f"CGPA: {fall_cgpa}", bold=True)
cgpa_row2[1].merge(cgpa_row2[3])
set_cell_left(cgpa_row2[1], "Scholastic Status: GAS", bold=True)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/Transcript_Table.docx"
doc.save(out)
print(f"Saved: {out}")
