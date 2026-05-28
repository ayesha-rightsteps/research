from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins (normal, not wide) ---
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# --- Default font throughout ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def heading1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def body_no_indent(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

# ============================================================
# TITLE BLOCK
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Research Proposal")
r.bold = True
r.font.size = Pt(15)
r.font.name = 'Times New Roman'
title.paragraph_format.space_after = Pt(2)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run(
    "Multi-UAV Coordination Using Deep Reinforcement Learning:\n"
    "Candidate Problem Areas and Proposed Direction"
)
r2.font.size = Pt(12)
r2.font.name = 'Times New Roman'
subtitle.paragraph_format.space_after = Pt(18)

# ============================================================
# 1. INTRODUCTION
# ============================================================
heading1(doc, "1. Introduction")

body(doc,
    "Autonomous multi-UAV systems have attracted considerable research attention over the past "
    "several years, driven largely by their potential applications in search and rescue, infrastructure "
    "inspection, and emergency communications. The common thread across this work is the use of "
    "deep reinforcement learning (DRL) to train UAVs to make decisions in complex, dynamic environments "
    "without requiring hand-engineered rules for every scenario."
)
body(doc,
    "After reviewing ten recent papers spanning path planning, swarm coordination, target assignment, "
    "collision avoidance, and LLM-integrated control, three research directions have emerged that are "
    "each genuinely open and within realistic scope for a masters-level project. This proposal describes "
    "all three, with the third being the primary recommendation. The reasoning behind the ordering "
    "is explained in each section."
)

# ============================================================
# 2. BACKGROUND: WHAT THE LITERATURE LEAVES OPEN
# ============================================================
heading1(doc, "2. Background: What the Literature Leaves Open")

body(doc,
    "The ten papers reviewed address different sub-problems of multi-UAV coordination. Papers focused "
    "on path planning (D3QN, DRL with 3D CNN) tend to work with a single drone. Papers that scale to "
    "larger swarms (mean-field DDPG) do so in two-dimensional environments with static obstacles and "
    "identical drones. Papers that tackle target assignment (TANet-TD3, DA-MAPPO) either limit testing "
    "to five drones or, in the case of DA-MAPPO, limit it to three. Papers that address collision "
    "avoidance (IGAT-MARL) do so in isolation, without any target assignment component. No paper "
    "combines all three requirements (target assignment, collision avoidance, and path planning) "
    "in a three-dimensional environment for a medium-sized swarm."
)
body(doc,
    "This gap is not incidental. Each paper's future work section points toward some version of "
    "what the others did not do. DA-MAPPO (dynamic target assignment) explicitly lists 3D environments "
    "and larger swarm sizes as next steps. IGAT-MARL (collision avoidance) lists the absence of target "
    "assignment as a limitation. The 3D path planning paper operates on a single drone and calls for "
    "multi-agent extension. The three candidate problems below are each drawn from these documented gaps."
)

# ============================================================
# 3. CANDIDATE PROBLEM A
# ============================================================
heading1(doc, "3. Candidate Problem A: Scaling Dynamic Target Assignment Beyond Three Drones")

body(doc,
    "DA-MAPPO, published in 2026 in the IEEE Internet of Things Journal, demonstrated that embedding "
    "real-time target assignment directly into each drone's observation vector substantially improves "
    "navigation performance in dynamic environments. The paper's ablation study showed that removing "
    "this assignment-augmented state caused success to drop to zero percent, confirming the mechanism's "
    "importance. However, every performance experiment in the paper uses exactly three drones navigating "
    "to three targets. The scalability of the approach beyond this configuration is untested."
)
body(doc,
    "A natural first extension would be to replicate the DA-MAPPO framework and evaluate it at five, "
    "eight, and ten drones while keeping the rest of the setup constant. The contribution here is "
    "empirical validation rather than architectural novelty: the question is whether the minimum-cost "
    "Hungarian assignment, when run at every decision step, remains stable and efficient as the number "
    "of drones and targets increases. The O(N³) complexity of the Hungarian algorithm means that at "
    "ten drones it is still fast, but at twenty it begins to raise concerns, and the paper does not "
    "address this."
)
body(doc,
    "This direction is implementable within a modest simulation setup and has a clear baseline. The "
    "limitation is that it is primarily a scaling study; the environment remains two-dimensional and "
    "obstacles remain static. If the results confirm that DA-MAPPO scales cleanly to ten drones, the "
    "contribution is meaningful but somewhat narrow. It is best suited as a project where the primary "
    "goal is gaining deep familiarity with MAPPO and multi-agent simulation before tackling a larger problem."
)

# ============================================================
# 4. CANDIDATE PROBLEM B
# ============================================================
heading1(doc, "4. Candidate Problem B: Dynamic Obstacle Handling in Multi-UAV Target Assignment")

body(doc,
    "Both DA-MAPPO and TANet-TD3 use static obstacles, meaning fixed walls or barriers that do not move "
    "during an episode. This is a significant simplification. In any real operational environment, "
    "a UAV swarm encounters other moving aircraft, vehicles, or unpredictable objects. The dynamic "
    "targets that DA-MAPPO handles are not obstacles; they are the goals the drones are trying to "
    "reach. Adding dynamic obstacles, which are objects the drones must avoid that move independently of "
    "the drones' actions, introduces a meaningfully different challenge."
)
body(doc,
    "The research question here would be: can an assignment-augmented MAPPO framework maintain high "
    "mission success when some fraction of obstacles in the environment are moving at varying speeds "
    "and in varying directions? The observation space would need to include the velocities of nearby "
    "dynamic obstacles in addition to the existing state information, and the reward would need to "
    "penalize proximity to moving obstacles appropriately."
)
body(doc,
    "This is a more substantive contribution than Problem A, and it remains achievable in simulation "
    "without introducing new architectural complexity. The main risk is that adding dynamic obstacles "
    "to a three-drone environment may still feel narrow in scope. Combining this with a modest scale "
    "increase (to five or six drones) would strengthen the contribution, but even then the environment "
    "remains two-dimensional. For that reason, this direction is recommended as a viable option but "
    "not the primary one."
)

# ============================================================
# 5. CANDIDATE PROBLEM C
# ============================================================
heading1(doc, "5. Candidate Problem C: 3D Multi-UAV Navigation with Integrated Target Assignment and Collision Avoidance")

body(doc,
    "This is the recommended direction. The motivation comes from a gap that is simultaneously "
    "documented by three separate papers, none of which fills it themselves. Paper 91 (DA-MAPPO) "
    "handles dynamic target assignment for three drones in 2D. Paper 9 (IGAT-MARL) handles collision "
    "avoidance for up to ten drones in 2D using a conflict-aware interaction graph. Paper 3 demonstrates "
    "effective DRL-based path planning in a genuine 3D environment, but only for a single drone. "
    "No existing work combines all three: real-time dynamic target assignment, collision avoidance, "
    "and three-dimensional navigation in a single multi-agent learning framework."
)

heading2(doc, "5.1 Problem Statement")

body(doc,
    "The proposed research addresses the following problem: given a swarm of five to ten homogeneous "
    "UAVs operating in a three-dimensional environment with static and dynamic obstacles, each drone "
    "must navigate to an assigned target while continuously avoiding collisions with its teammates "
    "and with obstacles, and the assignment of targets to drones must update in real time as the "
    "relative distances between drones and targets change. Each drone has only local observability: "
    "it can sense nearby obstacles via a limited-range sensor and communicate only with neighbors "
    "within a defined radius, and no central controller is available at execution time."
)
body(doc,
    "This is not an incremental restatement of prior work. DA-MAPPO cannot be directly applied here "
    "because it does not model 3D kinematics and does not include a collision avoidance mechanism "
    "beyond its basic reward penalty. IGAT-MARL cannot be directly applied because it has no target "
    "assignment component and treats all drones as undifferentiated agents flying toward undefined "
    "goals. Extending either one to cover the other's gap, and doing so in three dimensions, "
    "requires non-trivial design choices in observation space, reward structure, and the interaction "
    "between the assignment mechanism and the collision graph."
)

heading2(doc, "5.2 Proposed Approach")

body(doc,
    "The proposed framework builds on MAPPO (Multi-Agent Proximal Policy Optimization) as the "
    "core RL algorithm, chosen because it is stable under partial observability and has already "
    "been validated in the target assignment setting by DA-MAPPO. Three design components are "
    "proposed on top of the standard MAPPO setup."
)
body(doc,
    "First, assignment-augmented observations. At every decision step, a minimum-cost Hungarian "
    "assignment is computed globally based on current 3D Euclidean distances between drones and "
    "unoccupied targets. The assigned target's relative 3D position is appended directly to each "
    "drone's local observation vector, following the DA-MAPPO design. This gives each drone an "
    "explicit, always-current reference point for where it should be going, which DA-MAPPO's ablation "
    "study showed to be critical."
)
body(doc,
    "Second, conflict-aware interaction modeling. Inspired by IGAT-MARL, a dynamic interaction graph "
    "is maintained that connects only drone pairs predicted to be on a collision course within a "
    "defined time horizon. Each drone aggregates information from its current conflict neighbors "
    "using a graph attention mechanism before producing its navigation decision. This keeps the "
    "communication structure sparse and relevant rather than having every drone attend to every "
    "other drone at all times."
)
body(doc,
    "Third, a 3D action and observation space. Each drone's observation includes its own 3D position "
    "and velocity, the positions and velocities of neighbors within communication range, the relative "
    "position of its assigned target, and obstacle proximity readings in all six cardinal directions "
    "(up, down, north, south, east, west, or a denser ray-cast representation). Actions are continuous "
    "velocity commands in three dimensions, trained using MAPPO's clipped surrogate objective."
)
body(doc,
    "A curriculum training strategy is used, starting with five drones in a low-obstacle 3D "
    "environment with static targets, then introducing moving targets, then increasing obstacle "
    "density, and finally increasing swarm size to eight and ten drones. This follows the approach "
    "of DA-MAPPO and IGAT-MARL, both of which reported that progressive difficulty significantly "
    "improved training stability."
)

heading2(doc, "5.3 Evaluation Plan")

body(doc,
    "The primary evaluation metric is mission success rate: the fraction of episodes in which all "
    "drones reach their assigned targets without any inter-drone collision or obstacle collision "
    "within a fixed time limit. Secondary metrics include average trajectory length per drone, "
    "average time steps to completion, inter-drone collision count, and number of target reassignments "
    "per episode. The system will be evaluated across three environment configurations: thirty "
    "obstacles, forty obstacles, and fifty obstacles, matching the difficulty levels used in DA-MAPPO."
)
body(doc,
    "Baselines for comparison will include: (a) standard MAPPO without assignment augmentation, "
    "to quantify the contribution of the assignment mechanism; (b) the DA-MAPPO design ported to "
    "3D without the conflict-aware interaction graph, to quantify the contribution of the collision "
    "avoidance component; and (c) the IGAT-MARL approach ported to include a simple fixed assignment, "
    "to quantify the contribution of dynamic reassignment over static assignment. These three ablations "
    "directly mirror the claims being made about each component."
)

heading2(doc, "5.4 Why This Problem Over the Others")

body(doc,
    "Problem A (scaling DA-MAPPO) is essentially a subset of this proposal. Swarm scaling is already included "
    "in the experimental plan for Problem C and is embedded within a richer experimental context "
    "rather than being the sole contribution. Problem B (dynamic obstacles) is also subsumed here "
    "through the increasing obstacle density curriculum. The difference is that Problem C adds the "
    "3D dimension and the collision-avoidance mechanism, which together make the contribution "
    "substantially harder to dismiss as a straightforward extension of existing work."
)
body(doc,
    "From a publication standpoint, the three baselines in Problem C's evaluation plan directly "
    "correspond to the three papers whose gaps this work fills. A reviewer from any of those "
    "papers' venues, such as IEEE IoT Journal or Applied Soft Computing, would immediately "
    "recognize the experimental design as a direct and principled response to documented limitations "
    "in prior work. That alignment between the proposed contribution and the documented gaps in the "
    "literature is what makes a paper credible and reviewable."
)

# ============================================================
# 6. TIMELINE
# ============================================================
heading1(doc, "6. Proposed Timeline")

body_no_indent(doc, "The following schedule assumes approximately 15 to 18 months of active work.\n")

rows = [
    ("Months 1–3",   "Literature consolidation; 3D simulation environment setup in PyBullet or custom OpenAI Gym; implementation and validation of baseline MAPPO for three drones in 3D"),
    ("Months 4–6",   "Implementation of assignment-augmented observations and Hungarian assignment step; replication of DA-MAPPO results in 3D for three drones as sanity check"),
    ("Months 7–10",  "Integration of conflict-aware interaction graph (IGAT-inspired); scaling to five, eight, and ten drones with curriculum training; systematic evaluation against all baselines"),
    ("Months 11–13", "Ablation experiments; sensitivity analysis on obstacle density, swarm size, and communication radius; analysis of failure cases"),
    ("Months 14–18", "Writing; revision buffer; preparation for submission"),
]

table = doc.add_table(rows=len(rows)+1, cols=2)
table.style = 'Table Grid'

hdr = table.rows[0].cells
hdr[0].text = "Period"
hdr[1].text = "Activity"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

for i, (period, activity) in enumerate(rows):
    row = table.rows[i+1].cells
    row[0].text = period
    row[1].text = activity
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

doc.add_paragraph()  # spacing after table

# ============================================================
# 7. EXPECTED CONTRIBUTIONS
# ============================================================
heading1(doc, "7. Expected Contributions")

body(doc,
    "If the proposed work proceeds as outlined, the expected contributions are as follows. "
    "First, a working simulation framework for 3D multi-UAV navigation that can serve as a "
    "reproducible baseline for future work. This has value in itself given that most papers in this "
    "area use private or undocumented simulation setups. Second, an empirical answer to the question "
    "of whether real-time target reassignment and conflict-aware collision avoidance can be integrated "
    "in a single MAPPO policy without each component degrading the other. The interaction between "
    "these two mechanisms is not obvious and has not been studied. Third, a systematic comparison "
    "across five swarm sizes and three obstacle densities, which provides more comprehensive evidence "
    "than the existing papers in this space, most of which test only one or two swarm configurations."
)
body(doc,
    "The scope is intentionally constrained to homogeneous drones and simulated environments. "
    "Real hardware validation and heterogeneous swarm extension are acknowledged as future directions "
    "but are not part of the current scope. This constraint is intentional. Attempting to cover "
    "heterogeneous drones or real hardware within a masters project would either stretch the timeline "
    "unreasonably or force a shallower treatment of the core research question."
)

# ============================================================
# 8. REFERENCES (brief)
# ============================================================
heading1(doc, "8. Key References")

refs = [
    "Sheng et al. (2026). Dynamic Target Assignment and Cooperative Decision-Making for UAV Swarms Based on Multi-Agent Reinforcement Learning. IEEE Internet of Things Journal.",
    "Rezaee et al. (2026). Efficient Multi-Agent Deep Reinforcement Learning Algorithm for Multi UAV Collision Avoidance. Applied Soft Computing, Vol. 197.",
    "Fan et al. (2025). Dynamic Reward-Based Deep Reinforcement Learning Algorithm for UAV Path Planning in Large-Scale Environments. [Referenced as Paper 3 in review].",
    "Poudel & Moh (2026). MAML-Integrated Multi-Agent Reinforcement Learning for Adaptive Coalition-Based UAV Coordination in Disaster Scenarios. Internet of Things, Elsevier.",
    "Wang et al. (2025). RALLY: Role-Adaptive LLM-Driven Yoked Navigation for Agentic UAV Swarms. IEEE Open Journal of Vehicular Technology, Vol. 6.",
    "Xu et al. (2026). Scalable UAV Multi-Hop Networking via Multi-Agent Reinforcement Learning with Large Language Models. arXiv:2505.08448.",
]

for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ============================================================
# SAVE
# ============================================================
out_path = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/z. results/Research_Proposal.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
