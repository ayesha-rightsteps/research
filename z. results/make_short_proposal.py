from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

def b(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Masters Research Proposal")
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'
t.paragraph_format.space_after = Pt(4)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = st.add_run("3D Multi-UAV Navigation with Dynamic Target Assignment and Collision Avoidance")
r2.font.size = Pt(12)
r2.font.name = 'Times New Roman'
st.paragraph_format.space_after = Pt(16)

# Background
h(doc, "Background")
b(doc,
    "Recent work on multi-UAV coordination using deep reinforcement learning has made progress on "
    "individual sub-problems: dynamic target assignment (DA-MAPPO, 2026), graph-based collision "
    "avoidance (IGAT-MARL, 2026), and 3D single-drone path planning. However, no existing study "
    "combines these three in one framework. DA-MAPPO is limited to three drones in a 2D environment. "
    "IGAT-MARL handles collision avoidance only, with no assignment component. Both papers explicitly "
    "list the missing combination as future work."
)

# Problem
h(doc, "Research Problem")
b(doc,
    "How can a swarm of five to ten UAVs navigate to dynamically assigned targets in a 3D environment "
    "while avoiding collisions with each other and with obstacles, using only local observations and "
    "no central controller at execution time?"
)

# Approach
h(doc, "Proposed Approach")
b(doc,
    "The framework extends MAPPO with three additions. First, a minimum-cost Hungarian assignment "
    "runs at every decision step, and the assigned target's 3D position is fed directly into each "
    "drone's observation vector. Second, a conflict-aware interaction graph connects only drone pairs "
    "on a predicted collision course, keeping coordination sparse and relevant. Third, the action and "
    "observation spaces are extended to three dimensions with continuous velocity commands. Training "
    "follows a curriculum that begins with five drones and static targets, then adds moving targets, "
    "higher obstacle density, and larger swarm sizes up to ten drones."
)

# Baselines
h(doc, "Baselines for Comparison")
b(doc,
    "Results will be compared against: standard MAPPO without assignment augmentation, DA-MAPPO "
    "ported to 3D without the collision graph, and IGAT-MARL ported with a fixed assignment. "
    "Each ablation directly tests the contribution of one component."
)

# Timeline
h(doc, "Timeline")

rows = [
    ("Months 1-3",   "Environment setup; baseline MAPPO in 3D validated"),
    ("Months 4-7",   "Core framework implementation; curriculum training"),
    ("Months 8-11",  "Full experiments across swarm sizes and obstacle densities"),
    ("Months 12-15", "Writing and submission"),
]

table = doc.add_table(rows=len(rows) + 1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Period"
hdr[1].text = "Activity"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

for i, (period, activity) in enumerate(rows):
    row = table.rows[i + 1].cells
    row[0].text = period
    row[1].text = activity
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

doc.add_paragraph()

# Tools
h(doc, "Tools and Resources")
b(doc,
    "Python, PyTorch, PyBullet (simulation), MAPPO implementation. All tools are free and "
    "open source. No physical hardware required."
)

# Key refs
h(doc, "Key References")
refs = [
    "Sheng et al. (2026). Dynamic Target Assignment and Cooperative Decision-Making for UAV Swarms. IEEE IoT Journal.",
    "Rezaee et al. (2026). Efficient Multi-Agent DRL for Multi UAV Collision Avoidance. Applied Soft Computing.",
    "Fan et al. (2025). Dynamic Reward-Based DRL for UAV Path Planning in Large-Scale Environments.",
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/z. results/Research_Proposal_Short.docx"
doc.save(out)
print(f"Saved: {out}")
