from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page setup: landscape ────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Inches(13)
sec.page_height   = Inches(8.5)
sec.orientation   = 1
sec.top_margin    = Inches(0.5)
sec.bottom_margin = Inches(0.5)
sec.left_margin   = Inches(0.45)
sec.right_margin  = Inches(0.45)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(9)

# ── Helpers ──────────────────────────────────────────────────────────────────
def shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def write(cell, text, bold=False, center=False, color=None, size=9, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def section_row(table, label, color_hex, n_cols):
    row = table.add_row()
    for i in range(1, n_cols):
        row.cells[0].merge(row.cells[i])
    shade(row.cells[0], color_hex)
    write(row.cells[0], label, bold=True, center=True, color=(255,255,255), size=9)

# ── Colours ───────────────────────────────────────────────────────────────────
HDR_BG   = "1F3864"
SEC_SURV = "2E4057"
SEC_ALGO = "1B5E20"
SEC_UAV  = "6A1B9A"
SEC_REL  = "B71C1C"
ROW_A    = "FFFFFF"
ROW_B    = "F3F6FB"
GAP_BG   = "FFF8E1"

# ── Columns ───────────────────────────────────────────────────────────────────
HEADERS = [
    "Ref",
    "Author(s) & Year",
    "Key Contribution",
    "Method / Algorithm",
    "Setting\n(Env · Agents)",
    "Best Result",
    "Critical Gap\n(what it does NOT do)",
]
WIDTHS = [
    Inches(0.35),
    Inches(1.55),
    Inches(1.90),
    Inches(1.90),
    Inches(1.35),
    Inches(1.30),
    Inches(2.70),
]
N = len(HEADERS)

# ── Paper data ────────────────────────────────────────────────────────────────
# (ref, author_year, contribution, method, setting, result, gap)

SURVEYS = [
    ("[11]",
     "Govinda et al.\n(2025)",
     "Reviews DRL applications across transportation, robotics & UAV systems",
     "Systematic literature review",
     "Survey\nNo experiment",
     "Open gap identified: no unified framework for navigation + coordination",
     "Survey only — proposes no algorithm; does not address 3D multi-agent coordination"),

    ("[23]",
     "Aggarwal &\nKumar (2020)",
     "Taxonomises UAV path-planning from classical algorithms to DRL",
     "Systematic review of classical, evolutionary & DRL methods",
     "Survey\nIndoor & outdoor, 2D & 3D",
     "DRL identified as most promising for real-time dynamic environments",
     "Survey only — many reviewed methods now outdated; no multi-agent coverage"),

    ("[24]",
     "Oliehoek &\nAmato (2016)",
     "Provides the formal Dec-POMDP framework used by virtually all MARL UAV papers",
     "Theoretical formalization of Decentralized POMDPs",
     "Theoretical\nSmall illustrative examples",
     "Foundational problem formulation adopted across the entire MARL UAV field",
     "Exact solutions computationally intractable for large state spaces"),

    ("[25]",
     "Gronauer &\nDiepold (2022)",
     "Comprehensive taxonomy of MARL — identifies scalability & partial observability as the two open challenges",
     "Survey of MARL algorithms, environments & open problems",
     "Survey\nNo experiment",
     "Scalability & partial observability confirmed as dominant open challenges",
     "Survey only — post-2021 work not covered; no empirical contribution"),
]

ALGOS = [
    ("[12]",
     "Mnih et al.\n(2015)",
     "First end-to-end DRL agent reaching human-level play directly from raw pixels",
     "DQN: CNN + experience replay + target network",
     "Atari (49 games)\nSingle agent",
     "Human-level on 29/49 Atari games",
     "Overestimates Q-values; discrete actions only; no multi-agent extension"),

    ("[13]",
     "Van Hasselt et al.\n(2016)",
     "Fixes DQN's systematic Q-value overestimation using two decoupled networks",
     "Double DQN: decoupled action selection & evaluation",
     "Atari games\nSingle agent",
     "Reduces overestimation bias; outperforms DQN on most games",
     "Discrete actions only; does not improve sample efficiency"),

    ("[14]",
     "Wang et al.\n(2016)",
     "Separates value estimation into state-value and action-advantage streams",
     "Dueling Network: V(s) + A(s,a) via aggregation layer",
     "Atari games\nSingle agent",
     "Better generalisation when many actions have similar value",
     "Discrete actions only; no direct multi-agent extension"),

    ("[15]",
     "Schaul et al.\n(2016)",
     "Speeds learning by sampling the most informative transitions more often",
     "Prioritized Experience Replay (PER): TD-error based sampling",
     "Atari games\nSingle agent",
     "Significantly faster convergence from rare but important transitions",
     "Adds overhead for priority updates; sensitive to alpha/beta hyperparameters"),

    ("[16]",
     "Schulman et al.\n(2017)",
     "Stable on-policy training using a clipped objective — no trust-region math needed",
     "PPO: clipped surrogate objective, multiple minibatch epochs",
     "MuJoCo, Atari\nSingle agent",
     "Widely adopted standard baseline for continuous & discrete control",
     "On-policy; sample inefficient; slower than off-policy methods"),

    ("[17]",
     "Lowe et al.\n(2017)",
     "First CTDE actor-critic framework for mixed cooperative-competitive multi-agent tasks",
     "MADDPG: centralised critic + decentralised actor",
     "MPE\nMulti-agent",
     "Stable training across cooperative & competitive multi-agent scenarios",
     "Critic scales quadratically with agents; continuous actions only; struggles at large swarm sizes"),

    ("[18]",
     "Fujimoto et al.\n(2018)",
     "Eliminates overestimation in continuous-action actor-critic through three targeted fixes",
     "TD3: clipped double Q + delayed updates + target policy smoothing",
     "MuJoCo\nSingle agent",
     "State-of-the-art on continuous control; more stable than DDPG",
     "Off-policy; requires large replay buffer; not directly extendable to cooperative MARL"),

    ("[19]",
     "Rashid et al.\n(2018)",
     "Factorises joint Q-value into per-agent utilities while preserving global consistency",
     "QMIX: monotonic value mixing network, CTDE",
     "StarCraft SMAC\nMulti-agent",
     "Outperforms IQL and VDN on cooperative micromanagement tasks",
     "Monotonicity constraint limits representational capacity for non-monotonic interactions"),

    ("[20]",
     "Yang et al.\n(2018)",
     "Reduces multi-agent interaction complexity from O(N²) to O(N) via mean-field approximation",
     "Mean Field MARL: pairwise interactions replaced by mean-field term",
     "Battle game\nUp to 200 agents",
     "Scales to 200 agents where standard pairwise MARL fails computationally",
     "Loses individual agent information; accuracy degrades in heterogeneous or sparse swarms"),

    ("[21]",
     "Yu et al.\n(2022)",
     "Shows PPO with a shared centralised critic is surprisingly competitive with QMIX",
     "MAPPO: PPO + shared centralised critic + value normalisation",
     "SMAC, MPE\nMulti-agent",
     "Competitive with QMIX despite simpler design; widely reproduced baseline",
     "On-policy; sample inefficient; requires global state during training"),

    ("[22]",
     "Sunehag et al.\n(2018)",
     "Simplest value decomposition: joint Q-value as additive sum of per-agent Q-values",
     "VDN: additive Q-value factorisation, CTDE",
     "Cooperative matrix games\nMulti-agent",
     "Outperforms independent Q-learning; tractable and easy to implement",
     "Additivity assumption too restrictive for complex inter-agent dependencies; QMIX subsumes it"),
]

UAV = [
    ("[1]",
     "Tang et al.\n(2024)",
     "Single UAV navigates dynamic scenes using an improved DQN variant with PER",
     "Improved D3QN + Prioritized Experience Replay + heuristic target bias",
     "2D sim\n1 UAV",
     "95% success rate — outperforms A* and RRT on path quality",
     "Single drone only; 2D; discrete 8-direction actions; no multi-agent extension"),

    ("[2]",
     "Kong et al.\n(2024)",
     "Jointly solves target assignment and path planning for a small UAV team",
     "TANet-TD3: Target Assignment Network + TD3, Hungarian algorithm supervision",
     "2D sim\n5 UAVs",
     "Handles partial observability with dynamic targets; continuous actions",
     "Only 5 drones; 2D only; no collision-avoidance mechanism beyond reward"),

    ("[3]",
     "Jarray et al.\n(2025)",
     "Only paper to test DRL navigation in a true large-scale 3D environment",
     "DQN + 3D CNN feature extractor + dynamic step reward d3/(d1+d2)",
     "3D sim 25 km²\n1 UAV",
     "98% success in low obstacles; 85% in high-density obstacle scenes",
     "Single drone only; static obstacles; no team coordination possible"),

    ("[4]",
     "Zhang et al.\n(2025)",
     "Scales multi-UAV swarm coordination to 120 drones using mean-field approximation",
     "PO-WMFDDPG: mean-field DDPG + multi-head attention + CTDE",
     "2D sim\n20–120 UAVs",
     ">90% success at 120 drones where all standard MARL baselines collapse",
     "2D only; static obstacles; homogeneous drones; no dynamic target reassignment"),

    ("[5]",
     "Poudel & Moh\n(2026)",
     "Handles heterogeneous drones, drone failures & intermittent communication together",
     "RCTP: MAML + MA-DDPG + resource-aware coalition formation",
     "2D sim\n10–30 UAVs",
     "30–40% faster mission completion than baselines under failure scenarios",
     "2D only; up to 30 drones; MAML requires adaptation data at deployment"),

    ("[6]",
     "Xu et al.\n(2026)",
     "Uses GPT-4o knowledge distillation to initialise relay positions for multi-UAV networking",
     "MRLMN: IPPO + role-based reward + GPT-4o offline distillation",
     "2D sim\n12–24 UAVs",
     "+52% data rate and +27% user coverage versus baselines",
     "Networking task only; GPT-4o dependency; 2D horizontal movement only"),

    ("[7]",
     "Wang et al.\n(2025)",
     "LLM assigns drone roles then RL optimises coordination — generalises to unseen swarm sizes",
     "RALLY: LLM semantic consensus + RMIX role-value mixing",
     "MPE\n8–11 UAVs",
     "Zero-shot generalisation to swarm sizes not seen during training",
     "14.5 s inference latency; no real hardware validation; no target assignment"),

    ("[8]",
     "Khan et al.\n(2026)",
     "First fair large-scale comparison of LLM-based multi-agent coordination frameworks",
     "Systematic comparison: CrewAI vs AutoGen vs LangChain + hybrid routing",
     "CREW-WILDFIRE\n100+ agents",
     "96.1% success at 76% lower token cost than single-framework baselines",
     "Wildfire domain only; hybrid tested with only 6 agents; routing threshold not formally ablated"),
]

RELATED = [
    ("[9]",
     "Rezaee et al.\n(2026)",
     "Replaces dense all-to-all graph with sparse conflict-driven graph for collision avoidance",
     "IGAT-MARL: sparse conflict graph + Improved Graph Attention Network + curriculum",
     "BlueSky sim\n3–10 fixed-wing UAVs",
     "+17% reward; 10% fewer dangerous events; 44% fewer interaction edges vs. best baseline",
     "No target assignment — drones have no mission goal; fixed-wing only; 2D heading space"),

    ("[10]",
     "Sheng et al.\n(2026)",
     "Embeds real-time Hungarian assignment into each drone's observation at every decision step",
     "DA-MAPPO: per-step Hungarian + assignment-augmented observation + MAPPO + curriculum",
     "2D sim\n3 UAVs · 3 targets",
     "90–99% mission success; 25 pp above best baseline; ablation shows 0% without assignment",
     "2D only; 3 drones only; static obstacles; no collision-avoidance mechanism beyond reward penalty"),
]

# ── Build document ────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("Literature Review — Summary of Related Work")
r.bold = True; r.font.size = Pt(13); r.font.name = 'Times New Roman'
title_p.paragraph_format.space_after = Pt(4)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run(
    "25 papers grouped by contribution area. "
    "Read the Critical Gap column (amber) top to bottom — it tells the story of why this research is needed.")
r2.italic = True; r2.font.size = Pt(9); r2.font.name = 'Times New Roman'
sub_p.paragraph_format.space_after = Pt(8)

table = doc.add_table(rows=1, cols=N)
table.style = 'Table Grid'

for i, w in enumerate(WIDTHS):
    table.rows[0].cells[i].width = w

hdr_row = table.rows[0]
for i, h in enumerate(HEADERS):
    shade(hdr_row.cells[i], HDR_BG)
    write(hdr_row.cells[i], h, bold=True, center=True, color=(255,255,255), size=9)

def add_papers(papers, label, sec_color):
    section_row(table, label, sec_color, N)
    for idx, (ref, author, contrib, method, setting, result, gap) in enumerate(papers):
        row = table.add_row()
        bg  = ROW_A if idx % 2 == 0 else ROW_B
        shade(row.cells[0], bg);     write(row.cells[0], ref,     bold=True, center=True, size=9)
        shade(row.cells[1], bg);     write(row.cells[1], author,  size=9)
        shade(row.cells[2], bg);     write(row.cells[2], contrib, size=9)
        shade(row.cells[3], bg);     write(row.cells[3], method,  size=9)
        shade(row.cells[4], bg);     write(row.cells[4], setting, center=True, size=9)
        shade(row.cells[5], bg);     write(row.cells[5], result,  size=9)
        shade(row.cells[6], GAP_BG); write(row.cells[6], gap,     size=9, color=(100,60,0))
    # fix widths
    for row in table.rows:
        for i, w in enumerate(WIDTHS):
            row.cells[i].width = w

add_papers(SURVEYS, "GROUP 1 — SURVEYS & THEORETICAL FOUNDATIONS  (4 papers)", SEC_SURV)
add_papers(ALGOS,   "GROUP 2 — FOUNDATIONAL RL ALGORITHMS  (11 papers — building blocks used in UAV research)", SEC_ALGO)
add_papers(UAV,     "GROUP 3 — UAV COORDINATION PAPERS  (8 papers)", SEC_UAV)
add_papers(RELATED, "GROUP 4 — MOST RELATED WORK  (2 papers — directly motivates this research)", SEC_REL)

for row in table.rows:
    for i, w in enumerate(WIDTHS):
        row.cells[i].width = w

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(6)
rn = note_p.add_run(
    "Note: The Critical Gap column (amber) traces the open problems across all 25 papers. "
    "No reviewed paper combines dynamic target assignment with conflict-aware collision avoidance in a 3D multi-drone environment — "
    "this is the gap this research addresses.")
rn.italic = True; rn.font.size = Pt(8.5); rn.font.name = 'Times New Roman'
rn.font.color.rgb = RGBColor(100, 60, 0)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/Literature_Review_Table.docx"
doc.save(out)
print(f"Saved: {out}")
