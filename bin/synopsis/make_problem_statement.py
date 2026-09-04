from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def b(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def b_noi(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# ---- HEADER ----
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = meta.add_run("Ayesha Khalil\nSP25-RCS-009\nMS Computer Science")
r.font.size = Pt(11)
r.font.name = 'Times New Roman'
meta.paragraph_format.space_after = Pt(14)

# ---- TITLE ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Problem Statement")
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'
title.paragraph_format.space_after = Pt(6)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run(
    "Studying the Interference Between Target Assignment and Collision Avoidance "
    "in 3D Multi-UAV Coordination Using Multi-Agent Reinforcement Learning"
)
r2.font.size = Pt(12)
r2.font.name = 'Times New Roman'
subtitle.paragraph_format.space_after = Pt(18)

# ---- BACKGROUND ----
h(doc, "Background")

b(doc,
    "Coordinating a group of UAVs to reach different targets without colliding involves two "
    "objectives that are often in direct conflict: target assignment (which drone goes where) "
    "and collision avoidance (how drones stay out of each other's way). Most existing research "
    "solves one or the other. No published work has studied what happens when both are active "
    "in the same policy inside a 3D environment."
)

# ---- THE TWO PAPERS ----
h(doc, "The Two Papers That Define the Gap")

b(doc,
    "DA-MAPPO (Sheng et al., 2026) is the most advanced solution to target assignment in "
    "multi-UAV coordination. It runs the Hungarian algorithm at every decision step and feeds "
    "each drone's current assigned target directly into its observation. The paper's ablation "
    "proved how critical this is: removing the assignment information dropped mission success "
    "from over 90 percent to exactly zero. DA-MAPPO was tested with three drones in 2D and "
    "handled collisions only through a basic penalty term. The authors listed 3D environments "
    "and proper collision avoidance as future work."
)
b(doc,
    "IGAT-MARL (Rezaee et al., 2026) is the most advanced solution to multi-UAV collision "
    "avoidance. Rather than connecting every drone to every other drone, it builds a sparse "
    "graph that only links pairs currently on a predicted collision course. This reduced "
    "unnecessary interactions by 44 percent while improving safety. IGAT-MARL has no target "
    "assignment component at all. The authors listed task allocation as the obvious next step."
)

# ---- THE PROBLEM ----
h(doc, "The Problem")

b(doc,
    "When both mechanisms are placed inside the same policy in a 3D environment, they can "
    "directly contradict each other. The assignment mechanism tells a drone to fly toward "
    "its target. The conflict graph tells the same drone to deviate because another drone "
    "is on its path. In 2D, a sideways deviation keeps the drone roughly on course. In 3D, "
    "a drone climbing toward a target above it may be pushed horizontally by the conflict "
    "graph, taking it completely off its upward trajectory. Both instructions arrive at the "
    "same decision step and neither mechanism knows the other exists."
)
b(doc,
    "No existing paper has placed these two mechanisms in the same policy or tested whether "
    "DA-MAPPO's core mechanism, the assignment-augmented observation, stays effective when "
    "drones must also manage altitude and 3D collision avoidance simultaneously. The outcome "
    "is not known, and both papers confirm this gap in their own future work sections."
)

# ---- RESEARCH QUESTION ----
h(doc, "Research Question")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
p.paragraph_format.first_line_indent = Inches(0.3)
run1 = p.add_run(
    "When target assignment and collision avoidance operate within the same learned policy "
    "in a 3D multi-UAV environment, do the two mechanisms reinforce each other or interfere "
    "with each other, and under what conditions does each outcome occur?"
)
run1.font.name = 'Times New Roman'
run1.font.size = Pt(12)
run1.italic = True

# ---- SCOPE ----
h(doc, "Scope")

b(doc,
    "The study will use simulation (PyBullet) with swarm sizes of three, five, and eight "
    "drones. Targets will be dynamic. Real hardware, heterogeneous drones, and communication "
    "constraints are outside the scope of this work and are noted as natural future directions."
)

# ---- REFERENCES ----
h(doc, "References")

refs = [
    "Sheng, Y., Xie, X., Liu, H., & Li, J. (2026). Dynamic Target Assignment and Cooperative Decision-Making for UAV Swarms Based on Multi-Agent Reinforcement Learning. IEEE Internet of Things Journal.",
    "Rezaee, M. R., Abdul Hamid, N. A. W., Hussin, M., & Zukarnain, Z. A. (2026). Efficient Multi-Agent Deep Reinforcement Learning Algorithm for Multi UAV Collision Avoidance. Applied Soft Computing, Vol. 197.",
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/synopsis/ProblemStatement_Final.docx"
doc.save(out)
print(f"Saved: {out}")
