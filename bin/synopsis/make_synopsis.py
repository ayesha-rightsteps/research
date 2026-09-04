from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def b(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def b_noi(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# ---- HEADER ----
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = meta.add_run("Ayesha Khalil | SP25-RCS-009\nMS Computer Science")
r.font.size = Pt(11)
r.font.name = 'Times New Roman'
meta.paragraph_format.space_after = Pt(10)

# ---- TITLE ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run(
    "What Happens When a Drone Is Told to Go Forward and Avoid Collision at the Same Time:\n"
    "A Study of Target Assignment and Collision Avoidance Interference\n"
    "in 3D Multi-UAV Navigation"
)
r.bold = True
r.font.size = Pt(13)
r.font.name = 'Times New Roman'
title.paragraph_format.space_after = Pt(16)

# ---- 1. BACKGROUND ----
h(doc, "1. Background")

b(doc,
    "Getting a group of drones to reach different targets without hitting each other sounds "
    "straightforward, but it involves two objectives that pull in opposite directions. The "
    "first is target assignment: each drone needs to know which target is its responsibility "
    "and fly toward it. The second is collision avoidance: each drone needs to keep a safe "
    "distance from teammates, which sometimes means flying away from its target. In simple "
    "scenarios these two objectives coexist without much trouble. In dynamic environments "
    "with moving targets, changing assignments, and three-dimensional airspace, they can "
    "directly contradict each other."
)
b(doc,
    "Most existing research treats these as separate problems and solves them independently. "
    "The assignment paper does not worry about collisions in any serious way. The collision "
    "paper gives drones no targets to reach. Neither paper has to deal with what happens when "
    "both objectives are active simultaneously and one undermines the other."
)

# ---- 2. THE GAP ----
h(doc, "2. The Specific Gap")

b(doc,
    "Two papers published in 2026 represent the best available solutions to each half of this "
    "problem. DA-MAPPO (Sheng et al., 2026) introduced a real-time target assignment mechanism "
    "where the Hungarian algorithm runs at every decision step and the assigned target's position "
    "is embedded directly into each drone's observation. The paper's own ablation study confirmed "
    "how important this single mechanism is: when the assignment information was removed from the "
    "observation, mission success dropped from over 90 percent to exactly zero. The mechanism is "
    "not optional; it is the entire basis of the policy. The limitation is that DA-MAPPO was only "
    "tested with three drones in a two-dimensional environment, and collision avoidance was handled "
    "by nothing more than a penalty term in the reward function."
)
b(doc,
    "IGAT-MARL (Rezaee et al., 2026) took a fundamentally different approach to the collision "
    "problem. Instead of having every drone monitor every other drone, it builds a sparse graph "
    "that only connects pairs of drones currently predicted to be on a collision course. Drones "
    "not in conflict are ignored entirely. This reduced unnecessary interaction overhead by 44 "
    "percent while improving avoidance performance. The limitation is equally clear: IGAT-MARL "
    "has no target assignment component. The drones in that paper have no goals. They are simply "
    "avoiding each other with no mission to complete."
)
b(doc,
    "Both papers acknowledged this in their own writing. DA-MAPPO listed 3D environments and "
    "larger swarms as future work. IGAT-MARL listed target assignment as the obvious next step. "
    "Neither paper took the step the other described. The gap is not something I identified by "
    "guessing; it is something both sets of authors wrote down themselves."
)

# ---- 3. THE RESEARCH PROBLEM ----
h(doc, "3. The Research Problem")

b(doc,
    "The problem I want to study is not simply whether these two mechanisms can be combined. "
    "That would be an engineering task. The actual research question is what happens when they "
    "are combined in 3D space, specifically whether they work together or interfere with each other."
)
b(doc,
    "In two dimensions, a drone navigating toward a target and a drone avoiding a conflict are "
    "working in the same plane. A deviation to the left or right to avoid a collision still "
    "leaves the drone in roughly the right direction. In three dimensions this is no longer true. "
    "Consider a drone assigned to a target that is above and ahead. The most direct path is "
    "diagonal and upward. Now suppose another drone is directly above it on the same vertical "
    "path. The conflict graph identifies this pair as a collision risk and pushes the drone to "
    "deviate. But any horizontal deviation takes the drone off its upward path toward its target. "
    "The assignment mechanism says go up and forward. The collision mechanism says move sideways. "
    "They give contradictory instructions, and neither mechanism knows the other exists."
)
b(doc,
    "The question is whether a single learned policy, trained with both mechanisms active, can "
    "resolve this contradiction on its own through reinforcement learning, or whether the "
    "contradiction causes the policy to fail in ways that neither mechanism alone would produce. "
    "That question has not been studied, and the answer is not obvious."
)

# ---- 4. RESEARCH QUESTION ----
h(doc, "4. Research Question")

b_noi(doc,
    "Can a single MAPPO policy trained with both assignment-augmented observations and a "
    "conflict-aware interaction graph achieve above 85 percent mission success for five to "
    "eight UAVs navigating to dynamically assigned targets in a 3D environment, and does each "
    "mechanism contribute independently or do they reduce each other's effectiveness when "
    "active simultaneously?"
)

# ---- 5. OBJECTIVES ----
h(doc, "5. Objectives")

b(doc,
    "The first objective is to design a unified observation vector that carries both the "
    "assignment information from DA-MAPPO and the conflict neighborhood from IGAT-MARL, "
    "extended to three dimensions, and verify that each component still functions as the "
    "original papers reported when tested individually."
)
b(doc,
    "The second objective is to train the combined policy across swarm sizes of three, five, "
    "and eight drones and measure whether performance holds as the number of simultaneous "
    "assignment-conflict interactions grows."
)
b(doc,
    "The third objective is to run ablation experiments that isolate the contribution of each "
    "component, specifically to identify whether the two mechanisms cooperate, are neutral to "
    "each other, or actively interfere. The interference case is the most interesting result "
    "and would itself point toward a clearly defined follow-on problem."
)

# ---- 6. METHODOLOGY ----
h(doc, "6. Methodology")

b(doc,
    "The simulation will be built in PyBullet, a free physics-based environment that supports "
    "3D movement with realistic kinematic constraints. The core algorithm is MAPPO, selected "
    "because DA-MAPPO already validated it for this task type and it is stable under partial "
    "observability. At every decision step, each drone observes its own 3D position and velocity, "
    "the relative position of its currently assigned target, the positions of drones it is "
    "currently in conflict with, and proximity readings in six directions for obstacle awareness. "
    "The Hungarian assignment runs every step to keep the target allocation current. The conflict "
    "graph updates every step to keep avoidance focused only on genuinely dangerous pairs."
)
b(doc,
    "Training follows a four-stage curriculum. It begins with three drones and static targets "
    "to verify that the baseline replicates DA-MAPPO's reported results. It then moves to five "
    "drones with moving targets, then to eight drones with higher obstacle density, and finally "
    "to a mixed configuration to test generalization. The four baselines for comparison are: "
    "standard MAPPO with no assignment or conflict graph, DA-MAPPO ported to 3D with no conflict "
    "graph, IGAT-MARL with a fixed assignment and no dynamic reassignment, and the original "
    "DA-MAPPO in 2D with three drones to confirm the replication is accurate."
)

# ---- 7. EXPECTED OUTCOME ----
h(doc, "7. Expected Outcome")

b(doc,
    "The most straightforward outcome is that the combined policy works and each component "
    "contributes independently, which would validate the design and produce a working 3D "
    "multi-UAV coordination framework that does not currently exist in the literature. "
    "The more interesting outcome is that interference is observed, meaning the combined "
    "policy performs worse than one of its components alone under certain conditions. "
    "That result would identify exactly where and why the two mechanisms conflict, which is "
    "a concrete and specific contribution to the field regardless of the performance numbers. "
    "Either outcome is publishable because neither has been studied."
)

# ---- 8. TIMELINE ----
h(doc, "8. Timeline")

rows = [
    ("Months 1-3",   "PyBullet environment setup; replicate DA-MAPPO at 3 drones in 2D as validation"),
    ("Months 4-6",   "Extend to 3D; integrate conflict graph; train on 3 drones and verify both components work individually"),
    ("Months 7-10",  "Scale to 5 and 8 drones; full curriculum training; evaluate against all four baselines"),
    ("Months 11-13", "Ablation experiments; interference analysis; identify failure conditions"),
    ("Months 14-18", "Writing and submission"),
]

table = doc.add_table(rows=len(rows)+1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Period"
hdr[1].text = "Activity"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
for i, (period, activity) in enumerate(rows):
    row = table.rows[i+1].cells
    row[0].text = period
    row[1].text = activity
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

doc.add_paragraph()

# ---- 9. REFERENCES ----
h(doc, "References")

refs = [
    "Sheng, Y., Xie, X., Liu, H., & Li, J. (2026). Dynamic Target Assignment and Cooperative Decision-Making for UAV Swarms Based on Multi-Agent Reinforcement Learning. IEEE Internet of Things Journal.",
    "Rezaee, M. R., Abdul Hamid, N. A. W., Hussin, M., & Zukarnain, Z. A. (2026). Efficient Multi-Agent Deep Reinforcement Learning Algorithm for Multi UAV Collision Avoidance. Applied Soft Computing, Vol. 197.",
    "Fan, X. et al. (2025). Dynamic Reward-Based Deep Reinforcement Learning Algorithm for UAV Path Planning in Large-Scale Environments. Procedia Computer Science.",
    "Poudel, S. & Moh, S. (2026). MAML-Integrated Multi-Agent Reinforcement Learning for Adaptive Coalition-Based UAV Coordination in Disaster Scenarios. Internet of Things, Elsevier.",
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/synopsis/Synopsis_Final.docx"
doc.save(out)
print(f"Saved: {out}")
