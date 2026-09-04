from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

def b(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.first_line_indent = Inches(0.3)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

def b_noi(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

# ---- TITLE ----
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Masters Research Proposal")
r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'
t.paragraph_format.space_after = Pt(4)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = st.add_run(
    "Size-Agnostic Policy Learning for Multi-UAV Swarm Coordination:\n"
    "Training Once, Deploying at Any Scale"
)
r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
st.paragraph_format.space_after = Pt(18)

# ---- 1. THE CONCRETE PROBLEM ----
h(doc, "1. The Concrete Problem")

b(doc,
    "Every published multi-agent reinforcement learning (MARL) system for UAV coordination is trained "
    "and tested on a fixed swarm size. A policy trained for eight drones assumes eight drones. "
    "If a ninth drone joins, or if two drones fail mid-mission, the policy either crashes or "
    "produces undefined behavior because the input dimensionality no longer matches what the network "
    "was trained on. This is not a minor implementation detail. It is a fundamental design assumption "
    "baked into the architecture of every system in the current literature."
)
b(doc,
    "Consider a concrete scenario: a swarm of ten drones is deployed for a search mission. Midway "
    "through, three drones lose battery and drop out. The remaining seven must continue the mission. "
    "With any existing coordination policy, the team either falls back to independent flight with "
    "no coordination, or fails entirely because the learned policy was never exposed to a seven-drone "
    "configuration. Neither outcome is acceptable in practice."
)
b(doc,
    "This is the specific gap this research targets: the inability of current MARL policies to "
    "function at swarm sizes they were not explicitly trained on. The problem is not just about "
    "scaling up; it is about operating correctly across a range of sizes that changes dynamically "
    "during a mission, without retraining."
)

# ---- 2. WHY EXISTING METHODS CANNOT SOLVE THIS ----
h(doc, "2. Why Existing Methods Cannot Solve This")

b(doc,
    "Mean-field reinforcement learning, used in PO-WMFDDPG (2025), approximates all neighboring "
    "drones as a single average agent. This allows training at large fixed sizes (80 to 120 drones) "
    "but it does not address variable sizes. The mean-field approximation is still computed over a "
    "fixed expected population, and the policy is validated only at sizes close to the training size. "
    "More critically, mean-field loses individual neighbor information entirely, which matters for "
    "precise collision avoidance and targeted coordination."
)
b(doc,
    "RALLY (2025) showed partial generalization from eight drones to eleven without retraining, using "
    "an attention-based architecture. This is the closest existing result to what is proposed here. "
    "However, the generalization was observed as a side effect rather than studied as a research "
    "question, the range tested was narrow (three extra drones), and RALLY's primary contribution "
    "was LLM-based semantic reasoning, not size generalization. The mechanism behind the generalization "
    "was not analyzed, and it was not tested under drone failures during a mission."
)
b(doc,
    "DA-MAPPO (2026), the most recent target assignment paper, uses a fixed three-drone setup for all "
    "experiments. The paper explicitly limits its evaluation to this configuration and acknowledges "
    "scalability as an open problem. Its assignment mechanism relies on a global Hungarian algorithm "
    "that has O(N cubed) complexity, which does not scale, and its observation vector has a fixed "
    "length tied to the exact number of teammates."
)
b(doc,
    "No existing paper has set out to answer the question: what architectural and training choices "
    "allow a UAV coordination policy to perform acceptably across a 3x to 5x range of swarm sizes, "
    "including sizes smaller and larger than training, without any retraining at deployment time?"
)

# ---- 3. RESEARCH QUESTION ----
h(doc, "3. Research Question")

b_noi(doc,
    "Can a multi-UAV coordination policy, trained on randomly sampled swarm sizes between N=3 "
    "and N=10, maintain acceptable task completion performance when deployed at untrained swarm "
    "sizes between N=2 and N=20, including under mid-mission drone failures?"
)
b(doc,
    "This is a single, specific, falsifiable question. Either the policy generalizes or it does not. "
    "The answer depends on specific, testable architectural and training choices, described below."
)

# ---- 4. HYPOTHESIS ----
h(doc, "4. Hypothesis")

b(doc,
    "A permutation-invariant attention-based policy, trained with variable swarm sizes sampled "
    "randomly per episode, will generalize to unseen swarm sizes with less than 15 percent "
    "degradation in mission success rate across a range of N=2 to N=20 drones, and will recover "
    "within a fixed number of steps when drones fail mid-mission. Fixed-size policies and "
    "mean-field baselines will degrade significantly outside their training range."
)

# ---- 5. WHY THIS IS A NOVEL CONTRIBUTION ----
h(doc, "5. Why This Is a Novel Contribution")

b(doc,
    "The novelty is not in the attention mechanism itself, which exists, or in the coordination "
    "task, which has been studied. The novelty is in framing size-agnosticism as the primary "
    "objective and designing both the training procedure and the architecture around it. "
    "Specifically, three things distinguish this from prior work."
)
b(doc,
    "First, the training distribution is deliberately variable. Rather than training on a fixed N, "
    "each episode samples a swarm size uniformly from a range. The policy learns to coordinate "
    "regardless of how many agents are present, which forces the network to develop representations "
    "that do not depend on a specific number of neighbors."
)
b(doc,
    "Second, the observation design is neighbor-agnostic. Instead of concatenating the states of "
    "all N teammates (which produces a different-length vector for each N), the policy uses "
    "a graph attention network that processes neighbors as a set, with no fixed cardinality. "
    "This is architecturally compatible with any swarm size at inference time."
)
b(doc,
    "Third, mid-mission failure is treated as a training condition, not just a test condition. "
    "Some training episodes include drone dropouts at random timesteps, so the policy learns "
    "to redistribute tasks and re-coordinate after losses, not just at episode start."
)

# ---- 6. WHAT BREAKS WITHOUT THIS ----
h(doc, "6. The Specific Failure This Research Prevents")

b(doc,
    "Without size-agnostic training, a coordination policy fails in three observable ways when "
    "swarm size changes. First, if the input vector is too short (fewer drones than expected), "
    "the network receives garbage or zero-padded observations and produces unstable actions. "
    "Second, if the input vector is too long (more drones than expected), the policy simply "
    "cannot run. Third, even if the architecture can handle variable input (e.g., via padding), "
    "the policy has never learned to coordinate in that configuration and produces uncoordinated "
    "behavior equivalent to independent agents. Each of these failure modes will be demonstrated "
    "empirically in the experiments as part of the baseline analysis."
)

# ---- 7. METHODOLOGY ----
h(doc, "7. Methodology")

b(doc,
    "The simulation environment will be built in a standard multi-agent gym (e.g., PettingZoo or "
    "a custom PyBullet environment). Each episode, the swarm size N is sampled uniformly from "
    "a training range. Each drone observes its own state and the states of up to K nearest "
    "neighbors, regardless of total swarm size. A graph attention network aggregates neighbor "
    "information. The policy is trained using MAPPO with shared parameters across all drones, "
    "which naturally supports variable agent counts."
)
b(doc,
    "The task is cooperative navigation: each drone must reach a target while avoiding collisions "
    "with teammates and obstacles. Targets are assigned using a minimum-cost assignment at each "
    "step for swarms smaller than 12 drones, and a neighbor-local greedy assignment for larger "
    "swarms to maintain computational tractability. A fraction of training episodes include "
    "random drone failures after step 50, requiring remaining drones to re-assign and continue."
)

b(doc,
    "Evaluation compares four configurations: the proposed variable-N policy; a fixed-N MAPPO "
    "baseline trained at N=8; PO-WMFDDPG (mean-field baseline from Paper 4); and RALLY's "
    "attention-based policy retrained on this task. Each is tested at N=2, 4, 6, 8, 10, 14, "
    "18, and 20 drones. The primary metric is mission success rate at each N. Secondary metrics "
    "include post-failure recovery time and collision rate during recovery."
)

# ---- 8. TIMELINE ----
h(doc, "8. Timeline")

rows = [
    ("Months 1-3",   "Environment setup; implement fixed-N MAPPO baseline; confirm it breaks at off-training sizes"),
    ("Months 4-7",   "Design and train variable-N policy with graph attention; implement mid-mission failure training"),
    ("Months 8-11",  "Full evaluation across all N values; baseline comparisons; ablations on training range and failure rate"),
    ("Months 12-15", "Analysis of generalization mechanism; writing and submission"),
]

table = doc.add_table(rows=len(rows)+1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Period"
hdr[1].text = "Activity"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(11)
for i, (period, activity) in enumerate(rows):
    row = table.rows[i+1].cells
    row[0].text = period; row[1].text = activity
    for cell in row:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Times New Roman'; run.font.size = Pt(11)

doc.add_paragraph()

# ---- 9. EXPECTED CONTRIBUTION ----
h(doc, "9. Expected Contribution")

b(doc,
    "This work will produce, to the best of our knowledge, the first systematic study of "
    "size-agnostic policy learning for UAV swarm coordination. The contribution is not a new "
    "algorithm but a new training framework and a set of design principles for building "
    "coordination policies that remain functional when swarm composition changes at runtime. "
    "The work will also produce the first direct comparison of fixed-size policies, mean-field "
    "approximations, and attention-based policies specifically under off-training-distribution "
    "swarm sizes, which fills a gap that every scalable MARL paper in this domain has left open."
)

# ---- 10. REFERENCES ----
h(doc, "10. Key References")

refs = [
    "Wang et al. (2025). RALLY: Role-Adaptive LLM-Driven Yoked Navigation for Agentic UAV Swarms. IEEE OJVT, Vol. 6. [Closest related: shows partial size generalization as side effect]",
    "Sheng et al. (2026). Dynamic Target Assignment and Cooperative Decision-Making for UAV Swarms. IEEE IoT Journal. [Baseline: fixed 3-drone policy]",
    "PO-WMFDDPG paper (2025). Large-Scale UAV Swarm Path Planning Based on Mean-Field RL. [Baseline: fixed-size mean-field]",
    "Rezaee et al. (2026). Efficient Multi-Agent DRL for Multi UAV Collision Avoidance. Applied Soft Computing. [Baseline: fixed-size graph policy]",
    "Lowe et al. (2017). Multi-Agent Actor-Critic for Mixed Cooperative-Competitive Environments. NeurIPS. [MADDPG foundation]",
]
for ref in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'; run.font.size = Pt(11)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/z. results/Research_Proposal_Concrete.docx"
doc.save(out)
print(f"Saved: {out}")
