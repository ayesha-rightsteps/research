from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
section = doc.sections[0]
section.page_width    = Inches(13)
section.page_height   = Inches(8.5)
section.top_margin    = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin   = Inches(0.5)
section.right_margin  = Inches(0.5)
section.orientation   = 1

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(9)

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cw(cell, text, bold=False, center=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

papers = [
    # ----- Survey / review papers first -----
    {
        "ref":    "[11]",
        "authors":"Govinda, S., Brik, B., & Harous, S. (2025)",
        "method": "Systematic literature review of DRL applications across transportation, robotics, and UAV systems",
        "data":   "Survey of published DRL literature; no primary experimental data",
        "strength":"Identifies lack of unified frameworks combining navigation efficiency and inter-agent coordination as a key open gap",
        "limit":  "Survey only; no new algorithm proposed; limited coverage of multi-agent 3D coordination scenarios",
    },
    {
        "ref":    "[23]",
        "authors":"Aggarwal, S. & Kumar, N. (2020)",
        "method": "Systematic review of path planning algorithms for UAVs: classical, evolutionary, and DRL-based methods",
        "data":   "Review of published literature on UAV path planning across indoor, outdoor, 2D, and 3D environments",
        "strength":"Comprehensive taxonomy; identifies DRL as most promising for dynamic real-time environments; covers failure cases of classical methods",
        "limit":  "Survey only; no new algorithm; many reviewed methods predated modern MARL advances and are now outdated",
    },
    {
        "ref":    "[24]",
        "authors":"Oliehoek, F. A. & Amato, C. (2016)",
        "method": "Formal mathematical framework for Decentralized Partially Observable Markov Decision Processes (Dec-POMDPs)",
        "data":   "Theoretical formalization; small illustrative examples",
        "strength":"Foundational framework for multi-agent coordination under partial observability; provides rigorous problem formulation used by virtually all MARL UAV papers",
        "limit":  "Theoretical only; exact solutions are computationally intractable for large state spaces; scalability is an inherent structural limitation",
    },
    {
        "ref":    "[25]",
        "authors":"Gronauer, A. & Diepold, K. (2022)",
        "method": "Comprehensive survey and taxonomy of multi-agent deep reinforcement learning methods, algorithms, environments, and open challenges",
        "data":   "Survey of published MARL literature; no primary experimental data",
        "strength":"Identifies scalability and partial observability as the two major open challenges in MARL; provides complete taxonomy from independent learners to fully centralized methods",
        "limit":  "Survey only; field evolves rapidly so coverage of post-2021 work is limited; no empirical contribution",
    },
    # ----- Papers [1] to [22] in numerical order -----
    {
        "ref":    "[1]",
        "authors":"Tang, J., Liang, Y., & Li, K. (2024)",
        "method": "Improved D3QN: Dueling Double DQN + Prioritized Experience Replay + heuristic action bias toward target",
        "data":   "Custom 2D grid sim; single UAV; static & moving obstacles",
        "strength":"Handles moving obstacles; PER speeds convergence; outperforms A* and RRT on path quality",
        "limit":  "Single drone only; 2D; discrete 8-direction actions; no hardware validation",
    },
    {
        "ref":    "[2]",
        "authors":"Kong, X., Zhou, Y., Li, Z., & Wang, S. (2024)",
        "method": "TANet-TD3: Target Assignment Network + TD3; Hungarian algorithm supervision; POMDP formulation",
        "data":   "Custom 2D sim; 5 UAVs; moving targets; partial observability",
        "strength":"Jointly solves assignment + path planning; continuous actions; handles partial observability",
        "limit":  "Only 5 drones; 2D only; no scaling mechanism; no hardware test",
    },
    {
        "ref":    "[3]",
        "authors":"Jarray, R., Zaghbani, I., & Bouallègue, S. (2025)",
        "method": "DQN + 3D CNN feature extractor; dynamic reward formula d3/(d1+d2) per step; dual Q-networks",
        "data":   "Custom 3D sim; 25 km² environment; 4 obstacle-density scenarios; single UAV",
        "strength":"Only paper with genuine large-scale 3D environment; dense reward improves learning; beats PSO and GWO",
        "limit":  "Single drone; static obstacles; discrete actions; no multi-agent extension",
    },
    {
        "ref":    "[4]",
        "authors":"Zhang, Y. et al. (2025)",
        "method": "PO-WMFDDPG: DDPG + weighted mean-field + multi-head attention + CTDE; partial observability",
        "data":   "Custom 2D sim; swarm sizes 20–120 UAVs; static obstacles",
        "strength":"Scales to 120 drones; attention weighting is principled; >90% success where baselines collapse",
        "limit":  "2D only; static obstacles; homogeneous drones; ideal communication assumed",
    },
    {
        "ref":    "[5]",
        "authors":"Poudel, S. & Moh, S. (2026)",
        "method": "RCTP: MAML meta-learning + MA-DDPG + resource-aware coalition formation + LoS/NLoS communication modeling",
        "data":   "Custom disaster-response sim; 10–30 heterogeneous UAVs; dynamic obstacles; drone failure scenarios",
        "strength":"Handles heterogeneous drones, drone failures, intermittent communication, and energy constraints simultaneously; 30–40% faster missions",
        "limit":  "Only 10–30 drones; 2D; no comparison to mean-field methods; MAML needs adaptation data",
    },
    {
        "ref":    "[6]",
        "authors":"Xu, Y. et al. (2026)",
        "method": "MRLMN: IPPO + role-based reward decomposition + behavioral constraints + GPT-4o offline knowledge distillation via Hungarian matching",
        "data":   "Custom sim; 12–24 UAVs; 150 ground users; 3 base stations; 6.76–14.44 km² area",
        "strength":"52% higher data rate and 27% more user coverage than baselines; principled LLM-MARL integration; thorough ablation",
        "limit":  "Simulation only; proprietary GPT-4o dependency; no energy modeling; 2D horizontal movement only",
    },
    {
        "ref":    "[7]",
        "authors":"Wang, Z. et al. (2025)",
        "method": "RALLY: Two-stage LLM semantic consensus + RMIX role-value mixing (Commander/Coordinator/Executor); GPT-4o seeding; Qwen2.5-1.5B fine-tuned",
        "data":   "Multi-Agent Particle Environment (MPE); 8–11 drones; DS-CEFC task; one adversary drone",
        "strength":"Generalizes to unseen swarm sizes without retraining; theoretical proof; end-to-end deployment pipeline under 5 GB",
        "limit":  "14.48s inference latency; qualitative SITL validation only; no real hardware; no historical memory",
    },
    {
        "ref":    "[8]",
        "authors":"Khan, M. S., Khan, M. J., Sharif, M., & Yasmin, S. (2026)",
        "method": "Systematic comparison of CrewAI, AutoGen, LangChain + LangGraph-CrewAI hybrid with complexity-aware routing",
        "data":   "CREW-WILDFIRE benchmark; 100+ agents; 4 agent types; 17 task levels; 51 episodes per framework",
        "strength":"96.1% success at 76% lower token cost; first fair large-scale framework comparison; production-grade cost analysis",
        "limit":  "Wildfire domain only; hybrid validated with only 6 agents; routing threshold not formally ablated; 2.2s latency",
    },
    {
        "ref":    "[9]",
        "authors":"Rezaee, M. R., Abdul Hamid, N. A. W., Hussin, M., & Zukarnain, Z. A. (2026)",
        "method": "IGAT-MARL: Conflict-driven sparse interaction graph + Improved Graph Attention Network (stacked double-attention, residual connections) + curriculum + transfer learning",
        "data":   "BlueSky air traffic simulator (BADA); 3–10 fixed-wing UAVs; conflict-guaranteed episodes",
        "strength":"17% higher reward; 10% fewer dangerous separation events; 44% fewer interaction edges; comprehensive ablation",
        "limit":  "Fixed-wing only; 3 discrete heading actions; simulation only; no target assignment component",
    },
    {
        "ref":    "[10]",
        "authors":"Sheng, Y., Xie, X., Liu, H., & Li, J. (2026)",
        "method": "DA-MAPPO: Per-step Hungarian assignment + assignment-augmented observation vector + MAPPO + hierarchical 4-tier reward + curriculum training",
        "data":   "Custom sim; 3 UAVs; 3 dynamic targets; 30–50 static obstacles; 3 difficulty levels",
        "strength":"90–99% mission success; 25pp above best baseline; robust to 50% packet loss and 6-step delay; runs on Jetson Orin Nano",
        "limit":  "Only 3 drones tested; 2D; static obstacles; all-or-nothing success metric; no heterogeneous drones",
    },
    {
        "ref":    "[12]",
        "authors":"Mnih, V. et al. (2015)",
        "method": "DQN: CNN-based Q-network + experience replay + target network; trained end-to-end from raw pixel input",
        "data":   "49 Atari 2600 games (visual pixel input, 210x160 RGB frames)",
        "strength":"First end-to-end DRL from raw pixels; achieved human-level on 29/49 games; introduced experience replay and target network; foundational to all subsequent DQN variants",
        "limit":  "Overestimates action values (Q-value bias); discrete actions only; sample inefficient; single task per model",
    },
    {
        "ref":    "[13]",
        "authors":"Van Hasselt, H., Guez, A., & Silver, D. (2016)",
        "method": "Double DQN: decouples action selection and action evaluation using online and target networks to reduce overestimation bias",
        "data":   "Atari 2600 games (same benchmark as DQN)",
        "strength":"Reduces systematic Q-value overestimation; improves stability and final performance over DQN on most Atari games",
        "limit":  "Discrete action spaces only; does not address data efficiency or multi-step credit assignment",
    },
    {
        "ref":    "[14]",
        "authors":"Wang, Z. et al. (2016)",
        "method": "Dueling Network: separates state-value stream V(s) and advantage stream A(s,a); combined via aggregation layer",
        "data":   "Atari 2600 games",
        "strength":"Better policy evaluation in states where action choice matters less; improves generalization; compatible with any DQN variant",
        "limit":  "Discrete action space; benefit most visible when many actions have similar values; no direct multi-agent extension",
    },
    {
        "ref":    "[15]",
        "authors":"Schaul, T., Quan, J., Antonoglou, I., & Silver, D. (2016)",
        "method": "Prioritized Experience Replay (PER): samples replay buffer by TD-error priority; proportional and rank-based variants; importance sampling correction",
        "data":   "Atari 2600 games",
        "strength":"Learns faster from rare but informative transitions; significantly improves sample efficiency over uniform replay",
        "limit":  "Adds computational overhead for priority updates; sensitive to alpha and beta hyperparameters",
    },
    {
        "ref":    "[16]",
        "authors":"Schulman, J., Wolski, F., Dhariwal, A., Radford, A., & Klimov, O. (2017)",
        "method": "PPO: clipped surrogate objective to limit policy update magnitude; supports multiple minibatch epochs; no trust-region constraints needed",
        "data":   "MuJoCo continuous control; Atari games; robotics simulations",
        "strength":"Simple to implement; stable training; works for continuous and discrete actions; widely adopted as standard baseline",
        "limit":  "On-policy and sample inefficient; clip parameter requires tuning; slower convergence than off-policy methods in complex settings",
    },
    {
        "ref":    "[17]",
        "authors":"Lowe, R. et al. (2017)",
        "method": "MADDPG: CTDE with centralized critic on joint state and decentralized actor using local observations only; extends DDPG to multi-agent",
        "data":   "Multi-Agent Particle Environment (MPE); cooperative, competitive, and mixed scenarios",
        "strength":"Enables stable multi-agent training; handles cooperative and competitive tasks; decentralized execution",
        "limit":  "Critic input size grows with number of agents; struggles with large swarms; limited to continuous action spaces",
    },
    {
        "ref":    "[18]",
        "authors":"Fujimoto, S., Hoof, H., & Meger, D. (2018)",
        "method": "TD3: Twin Delayed DDPG with clipped double Q-learning + delayed policy updates + target policy smoothing to reduce overestimation in actor-critic",
        "data":   "MuJoCo continuous control benchmarks (HalfCheetah, Hopper, Walker2D, Ant)",
        "strength":"Reduces overestimation in actor-critic methods; more stable than DDPG; state-of-the-art on continuous control tasks",
        "limit":  "Off-policy; requires large replay buffer; harder to extend to cooperative multi-agent settings directly",
    },
    {
        "ref":    "[19]",
        "authors":"Rashid, T. et al. (2018)",
        "method": "QMIX: monotonic mixing network factorizes joint Q-value into per-agent utilities; CTDE; enforces Individual-Global-Max (IGM) property",
        "data":   "StarCraft Multi-Agent Challenge (SMAC); cooperative unit micromanagement tasks",
        "strength":"Handles cooperative multi-agent tasks well; monotonicity ensures consistent joint action selection; outperforms IQL and VDN",
        "limit":  "Monotonicity constraint limits representational capacity; cannot model non-monotonic agent interactions",
    },
    {
        "ref":    "[20]",
        "authors":"Yang, Y. et al. (2018)",
        "method": "Mean Field MARL: approximates interactions between one agent and all others as interaction with a single mean-field agent; reduces pairwise complexity from O(N^2) to O(N)",
        "data":   "Ising model; network routing; battle game with up to 200 agents",
        "strength":"Scales to large agent populations; linear complexity; enables coordination where pairwise MARL fails computationally",
        "limit":  "Mean-field loses individual agent information; accuracy degrades in heterogeneous or spatially sparse distributions",
    },
    {
        "ref":    "[21]",
        "authors":"Yu, C. et al. (2022)",
        "method": "MAPPO: PPO adapted for cooperative MARL with shared centralized critic on global state; value normalization; clipped updates",
        "data":   "StarCraft Multi-Agent Challenge (SMAC); Multi-Agent Particle Environment (MPE)",
        "strength":"Surprisingly competitive with QMIX despite simplicity; stable training; reproducible; adopted as standard MARL baseline in recent UAV work",
        "limit":  "On-policy and sample inefficient; requires global state for critic during training; performance gap narrows with larger teams",
    },
    {
        "ref":    "[22]",
        "authors":"Sunehag, P. et al. (2018)",
        "method": "VDN: Value Decomposition Networks — joint Q-value decomposed as additive sum of per-agent Q-values; CTDE; simplest value decomposition",
        "data":   "Decentralized cooperative matrix games; StarCraft-style cooperative tasks",
        "strength":"Simple and tractable; enables cooperative training with decentralized execution; outperforms independent Q-learning",
        "limit":  "Additivity assumption is too restrictive for complex inter-agent dependencies; QMIX generalizes VDN and outperforms it",
    },
]

# ---------------------------------------------------------------
HEADER_BG = "1F3864"
ROW_A_BG  = "FFFFFF"
ROW_B_BG  = "EEF2FA"
cols      = ["Ref #", "Author(s) & Year", "Methodology", "Dataset / Samples", "Strength", "Limitation"]
col_w     = [Inches(0.42), Inches(1.6), Inches(2.4), Inches(1.65), Inches(2.1), Inches(2.1)]

table = doc.add_table(rows=1 + len(papers), cols=len(cols))
table.style = 'Table Grid'

for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_w[i]

# header
for i, label in enumerate(cols):
    c = table.rows[0].cells[i]
    set_bg(c, HEADER_BG)
    cw(c, label, bold=True, center=True, size=9.5, color=(255, 255, 255))

# data rows
for idx, p in enumerate(papers):
    row  = table.rows[idx + 1]
    bg   = ROW_A_BG if idx % 2 == 0 else ROW_B_BG
    for cell in row.cells:
        set_bg(cell, bg)
    cw(row.cells[0], p["ref"],     bold=True, center=True)
    cw(row.cells[1], p["authors"])
    cw(row.cells[2], p["method"])
    cw(row.cells[3], p["data"])
    cw(row.cells[4], p["strength"])
    cw(row.cells[5], p["limit"])

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/Literature_Review_Table_25.docx"
doc.save(out)
print(f"Saved: {out}")
