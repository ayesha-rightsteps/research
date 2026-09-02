from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sec = doc.sections[0]
sec.page_width    = Inches(11.5)
sec.page_height   = Inches(8.5)
sec.orientation   = 1
sec.top_margin    = Inches(0.55)
sec.bottom_margin = Inches(0.55)
sec.left_margin   = Inches(0.5)
sec.right_margin  = Inches(0.5)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(10)

def shade(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cw(cell, text, bold=False, center=False, color=None, size=9, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def sec_row(table, label, color_hex, n):
    row = table.add_row()
    for i in range(1, n):
        row.cells[0].merge(row.cells[i])
    shade(row.cells[0], color_hex)
    cw(row.cells[0], label, bold=True, center=True, color=(255,255,255), size=9)

HDR = "1F3864"
S1  = "2C3E6B"
S2  = "1B5E20"
S3  = "4A1486"
S4  = "7B0000"
RA  = "FFFFFF"
RB  = "EEF2FA"
GAP = "FFF8E1"

HEADS  = ["Ref", "Author(s) & Year", "Method", "Key Result", "Gap"]
WIDTHS = [Inches(0.38), Inches(1.45), Inches(2.60), Inches(2.10), Inches(3.85)]
N = 5

# (ref, author, method, result, gap)
SURVEYS = [
    ("[11]", "Govinda et al.\n(2025)",
     "Systematic review of DRL in UAV systems",
     "No unified framework for navigation + coordination",
     "Survey only; no algorithm proposed"),

    ("[23]", "Aggarwal &\nKumar (2020)",
     "Review of UAV path-planning algorithms",
     "DRL identified as best for dynamic environments",
     "Survey only; no multi-agent coverage"),

    ("[24]", "Oliehoek &\nAmato (2016)",
     "Formal Dec-POMDP framework for multi-agent coordination",
     "Standard formulation used by all MARL UAV papers",
     "Theoretical only; intractable for large state spaces"),

    ("[25]", "Gronauer &\nDiepold (2022)",
     "Survey and taxonomy of MARL methods",
     "Scalability and partial observability as top open challenges",
     "Survey only; no empirical contribution"),
]

ALGOS = [
    ("[12]", "Mnih et al.\n(2015)",
     "DQN: CNN + experience replay + target network",
     "Human-level on 29/49 Atari games",
     "Discrete actions only; no multi-agent extension"),

    ("[13]", "Van Hasselt et al.\n(2016)",
     "Double DQN: decoupled action selection and evaluation",
     "Reduces Q-value overestimation over DQN",
     "Discrete actions only"),

    ("[14]", "Wang et al.\n(2016)",
     "Dueling Network: V(s) + A(s,a) streams",
     "Better generalisation when action choice matters less",
     "Discrete actions only; no multi-agent extension"),

    ("[15]", "Schaul et al.\n(2016)",
     "Prioritized Experience Replay (PER): TD-error sampling",
     "Faster convergence from rare informative transitions",
     "Adds overhead; sensitive to hyperparameters"),

    ("[16]", "Schulman et al.\n(2017)",
     "PPO: clipped surrogate objective",
     "Stable baseline for continuous and discrete control",
     "On-policy; sample inefficient"),

    ("[17]", "Lowe et al.\n(2017)",
     "MADDPG: centralised critic + decentralised actor",
     "Stable training across cooperative and competitive tasks",
     "Scales poorly with number of agents"),

    ("[18]", "Fujimoto et al.\n(2018)",
     "TD3: clipped double Q + delayed updates + target smoothing",
     "State-of-the-art on continuous control benchmarks",
     "Not directly extendable to cooperative MARL"),

    ("[19]", "Rashid et al.\n(2018)",
     "QMIX: monotonic joint Q-value factorisation",
     "Outperforms IQL and VDN on cooperative tasks",
     "Monotonicity limits representational capacity"),

    ("[20]", "Yang et al.\n(2018)",
     "Mean Field MARL: pairwise → mean-field interaction",
     "Scales to 200 agents with linear complexity",
     "Loses individual agent detail in heterogeneous swarms"),

    ("[21]", "Yu et al.\n(2022)",
     "MAPPO: PPO + shared centralised critic",
     "Competitive with QMIX; widely adopted MARL baseline",
     "On-policy; requires global state during training"),

    ("[22]", "Sunehag et al.\n(2018)",
     "VDN: additive per-agent Q-value decomposition",
     "Outperforms independent Q-learning; simple and tractable",
     "Too restrictive for complex inter-agent dependencies"),
]

UAV = [
    ("[1]", "Tang et al.\n(2024)",
     "Improved D3QN + PER + heuristic target bias",
     "95% success; outperforms A* and RRT",
     "Single UAV; 2D only"),

    ("[2]", "Kong et al.\n(2024)",
     "TANet-TD3: assignment + path planning, Hungarian supervision",
     "Handles dynamic targets with partial observability",
     "5 drones only; 2D only; no collision avoidance"),

    ("[3]", "Jarray et al.\n(2025)",
     "DQN + 3D CNN + dynamic step reward",
     "98% success in low-obstacle 3D environment (25 km²)",
     "Single UAV; static obstacles only"),

    ("[4]", "Zhang et al.\n(2025)",
     "PO-WMFDDPG: mean-field DDPG + multi-head attention",
     ">90% success at 120 UAVs where baselines collapse",
     "2D only; static obstacles; no dynamic reassignment"),

    ("[5]", "Poudel & Moh\n(2026)",
     "RCTP: MAML + MA-DDPG + coalition formation",
     "30-40% faster mission completion under drone failures",
     "2D only; up to 30 drones"),

    ("[6]", "Xu et al.\n(2026)",
     "MRLMN: IPPO + role-based reward + GPT-4o distillation",
     "+52% data rate; +27% user coverage",
     "Networking task only; 2D; GPT-4o dependency"),

    ("[7]", "Wang et al.\n(2025)",
     "RALLY: LLM role assignment + RMIX value mixing",
     "Generalises to unseen swarm sizes without retraining",
     "14.5s latency; no target assignment"),

    ("[8]", "Khan et al.\n(2026)",
     "Comparison: CrewAI vs AutoGen vs LangChain + hybrid",
     "96.1% success at 76% lower token cost",
     "Wildfire domain only; hybrid tested with 6 agents only"),
]

RELATED = [
    ("[9]", "Rezaee et al.\n(2026)",
     "IGAT-MARL: sparse conflict graph + Graph Attention Network",
     "+17% reward; 44% fewer graph edges vs. best baseline",
     "No target assignment; fixed-wing only; 2D only"),

    ("[10]", "Sheng et al.\n(2026)",
     "DA-MAPPO: per-step Hungarian assignment in observation + MAPPO",
     "90-99% mission success; 25 pp above best baseline",
     "2D only; 3 drones only; no collision-avoidance mechanism"),
]

# build
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title_p.add_run("Table 1. Summary of Related Work")
rt.bold = True; rt.font.size = Pt(11); rt.font.name = 'Times New Roman'
title_p.paragraph_format.space_after = Pt(5)

table = doc.add_table(rows=1, cols=N)
table.style = 'Table Grid'

for i, w in enumerate(WIDTHS):
    table.rows[0].cells[i].width = w

hrow = table.rows[0]
for i, h in enumerate(HEADS):
    shade(hrow.cells[i], HDR)
    cw(hrow.cells[i], h, bold=True, center=True, color=(255,255,255), size=9.5)

def add_group(papers, label, scol):
    sec_row(table, label, scol, N)
    for idx, (ref, author, method, result, gap) in enumerate(papers):
        row = table.add_row()
        bg  = RA if idx % 2 == 0 else RB
        shade(row.cells[0], bg);  cw(row.cells[0], ref,    bold=True, center=True)
        shade(row.cells[1], bg);  cw(row.cells[1], author)
        shade(row.cells[2], bg);  cw(row.cells[2], method)
        shade(row.cells[3], bg);  cw(row.cells[3], result)
        shade(row.cells[4], GAP); cw(row.cells[4], gap, color=(100,60,0))
    for row in table.rows:
        for i, w in enumerate(WIDTHS):
            row.cells[i].width = w

add_group(SURVEYS, "SURVEYS & THEORETICAL FOUNDATIONS", S1)
add_group(ALGOS,   "FOUNDATIONAL RL ALGORITHMS", S2)
add_group(UAV,     "UAV COORDINATION PAPERS", S3)
add_group(RELATED, "MOST RELATED WORK", S4)

for row in table.rows:
    for i, w in enumerate(WIDTHS):
        row.cells[i].width = w

note_p = doc.add_paragraph()
note_p.paragraph_format.space_before = Pt(5)
rn = note_p.add_run(
    "Note: Gap column (highlighted) — no existing work integrates dynamic target assignment "
    "with conflict-aware collision avoidance in a single learned policy.")
rn.italic = True; rn.font.size = Pt(9)
rn.font.name = 'Times New Roman'
rn.font.color.rgb = RGBColor(100, 60, 0)

out = "/Users/rightsteps/Masters/deep reinforcement learning/ayesha/newwwwww_lasttttttt/Literature_Review_Table.docx"
doc.save(out)
print(f"Saved: {out}")
